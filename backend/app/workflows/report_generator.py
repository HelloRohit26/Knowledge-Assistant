"""
Report Generator — Exporting Reports to PDF and DOCX
"""

import io
import re
from typing import Dict, Any, Optional
from app.core.logger import logger

# Try ReportLab for PDF generation
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

# Try python-docx for DOCX generation
try:
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def generate_pdf_report(title: str, category: str, markdown_content: str, metadata: Optional[Dict[str, Any]] = None) -> bytes:
    """
    Generates a PDF file from Markdown report content.
    """
    if REPORTLAB_AVAILABLE:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=54,
            leftMargin=54,
            topMargin=54,
            bottomMargin=54
        )

        styles = getSampleStyleSheet()

        # Custom Enterprise Styling
        title_style = ParagraphStyle(
            'DocTitle',
            parent=styles['Heading1'],
            fontSize=22,
            leading=26,
            textColor=colors.HexColor('#0F172A'),
            fontName='Helvetica-Bold',
            spaceAfter=6
        )

        meta_style = ParagraphStyle(
            'DocMeta',
            parent=styles['Normal'],
            fontSize=9,
            leading=12,
            textColor=colors.HexColor('#0284C7'),
            fontName='Helvetica-Bold',
            spaceAfter=15
        )

        h2_style = ParagraphStyle(
            'DocH2',
            parent=styles['Heading2'],
            fontSize=14,
            leading=18,
            textColor=colors.HexColor('#1E293B'),
            fontName='Helvetica-Bold',
            spaceBefore=14,
            spaceAfter=6
        )

        h3_style = ParagraphStyle(
            'DocH3',
            parent=styles['Heading3'],
            fontSize=11,
            leading=15,
            textColor=colors.HexColor('#334155'),
            fontName='Helvetica-Bold',
            spaceBefore=10,
            spaceAfter=4
        )

        body_style = ParagraphStyle(
            'DocBody',
            parent=styles['Normal'],
            fontSize=10,
            leading=14,
            textColor=colors.HexColor('#334155'),
            fontName='Helvetica',
            spaceAfter=6
        )

        bullet_style = ParagraphStyle(
            'DocBullet',
            parent=styles['Normal'],
            fontSize=10,
            leading=14,
            textColor=colors.HexColor('#334155'),
            fontName='Helvetica',
            leftIndent=15,
            spaceAfter=3
        )

        quote_style = ParagraphStyle(
            'DocQuote',
            parent=styles['Normal'],
            fontSize=9.5,
            leading=13.5,
            textColor=colors.HexColor('#475569'),
            fontName='Helvetica-Oblique',
            leftIndent=12,
            spaceAfter=6
        )

        story = []

        # Header Title
        story.append(Paragraph(_safe_reportlab_text(title), title_style))
        story.append(Paragraph(f"DEPARTMENT: {_safe_reportlab_text(category.upper())}  |  ENTERPRISE AI KNOWLEDGE INTELLIGENCE PLATFORM", meta_style))
        story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor('#0284C7'), spaceAfter=15))

        # Parse Markdown lines simply
        lines = markdown_content.split('\n')
        for line in lines:
            line_str = line.strip()
            if not line_str:
                continue

            if line_str.startswith('# '):
                # Skip main title if duplicate
                if line_str.replace('# ', '').strip().lower() != title.lower():
                    story.append(Paragraph(_safe_reportlab_text(line_str.replace('# ', '')), h2_style))
            elif line_str.startswith('## '):
                story.append(Paragraph(_safe_reportlab_text(line_str.replace('## ', '')), h2_style))
            elif line_str.startswith('### '):
                story.append(Paragraph(_safe_reportlab_text(line_str.replace('### ', '')), h3_style))
            elif line_str.startswith('- ') or line_str.startswith('* '):
                clean_text = _safe_reportlab_text(line_str[2:])
                story.append(Paragraph(f"• {clean_text}", bullet_style))
            elif re.match(r'^\d+\.\s', line_str):
                clean_text = _safe_reportlab_text(re.sub(r'^\d+\.\s', '', line_str))
                story.append(Paragraph(f"{line_str.split('.')[0]}. {clean_text}", bullet_style))
            elif line_str.startswith('>') or line_str.startswith('|'):
                clean_text = _safe_reportlab_text(line_str.replace('>', '').replace('|', ' '))
                story.append(Paragraph(clean_text, quote_style))
            else:
                clean_text = _safe_reportlab_text(line_str)
                story.append(Paragraph(clean_text, body_style))

        try:
            doc.build(story)
            pdf_bytes = buffer.getvalue()
            buffer.close()
            return pdf_bytes
        except Exception as err:
            logger.warning(f"ReportLab XML build failed ({err}), falling back to plain-text PDF builder.")
            return _generate_plain_pdf_fallback(title, category, markdown_content)

    else:
        # Fallback text buffer formatted as PDF header
        buffer = io.BytesIO()
        header = f"%PDF-1.4\n1 0 obj << /Title ({title}) /Subject ({category}) >>\n"
        buffer.write(header.encode('utf-8'))
        buffer.write(markdown_content.encode('utf-8'))
        return buffer.getvalue()


def generate_docx_report(title: str, category: str, markdown_content: str, metadata: Optional[Dict[str, Any]] = None) -> bytes:
    """
    Generates a DOCX file from Markdown report content.
    """
    if DOCX_AVAILABLE:
        doc = docx.Document()

        # Document Header Title
        h1 = doc.add_heading(title, level=0)
        h1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in h1.runs:
            run.font.color.rgb = RGBColor(15, 23, 42)
            run.font.name = 'Arial'

        # Metadata Subtitle
        sub = doc.add_paragraph(f"DEPARTMENT: {category.upper()}  |  ENTERPRISE KNOWLEDGE INTELLIGENCE PLATFORM")
        for run in sub.runs:
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.color.rgb = RGBColor(2, 132, 199)

        doc.add_paragraph("=" * 60)

        # Parse Markdown lines
        lines = markdown_content.split('\n')
        for line in lines:
            line_str = line.strip()
            if not line_str:
                continue

            if line_str.startswith('## '):
                doc.add_heading(line_str.replace('## ', ''), level=1)
            elif line_str.startswith('### '):
                doc.add_heading(line_str.replace('### ', ''), level=2)
            elif line_str.startswith('- ') or line_str.startswith('* '):
                doc.add_paragraph(_clean_markdown_inline(line_str[2:]), style='List Bullet')
            elif re.match(r'^\d+\.\s', line_str):
                doc.add_paragraph(_clean_markdown_inline(re.sub(r'^\d+\.\s', '', line_str)), style='List Number')
            else:
                if not line_str.startswith('# '):
                    doc.add_paragraph(_clean_markdown_inline(line_str))

        buffer = io.BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()
        buffer.close()
        return docx_bytes

    else:
        # Fallback formatted text buffer for docx download
        buffer = io.BytesIO()
        content = f"Title: {title}\nDepartment: {category}\n\n{markdown_content}"
        buffer.write(content.encode('utf-8'))
        return buffer.getvalue()


def _clean_markdown_inline(text: str) -> str:
    """
    Strips raw markdown syntax like **, *, `, etc. for plain text PDF/DOCX rendering.
    """
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    text = re.sub(r'`(.*?)`', r'\1', text)
    return text


def _safe_reportlab_text(text: str) -> str:
    """
    Escapes special XML characters and formats bold/italic safely for ReportLab Paragraphs.
    """
    # Escape ampersands and angle brackets first
    text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

    # Replace markdown **bold** with <b>bold</b>
    text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
    # Replace markdown *italic* with <i>italic</i>
    text = re.sub(r'\*(.*?)\*', r'<i>\1</i>', text)
    # Strip backticks
    text = text.replace('`', '')
    return text


def _generate_plain_pdf_fallback(title: str, category: str, markdown_content: str) -> bytes:
    """
    Fallback plain-text PDF builder when ReportLab XML tag parsing encounters invalid markup.
    """
    try:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54)
        styles = getSampleStyleSheet()
        plain_style = ParagraphStyle('PlainStyle', parent=styles['Normal'], fontSize=9.5, leading=13.5, textColor=colors.HexColor('#1E293B'))

        story = [
            Paragraph(f"<b>{title}</b>", styles['Heading1']),
            Paragraph(f"DEPARTMENT: {category.upper()}", styles['Normal']),
            Spacer(1, 12)
        ]

        plain_text = _clean_markdown_inline(markdown_content)
        for paragraph_block in plain_text.split('\n'):
            if paragraph_block.strip():
                # Escape XML chars safely
                safe_p = paragraph_block.strip().replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                story.append(Paragraph(safe_p, plain_style))
                story.append(Spacer(1, 4))

        doc.build(story)
        pdf_bytes = buffer.getvalue()
        buffer.close()
        return pdf_bytes
    except Exception as e:
        logger.error(f"Plain PDF fallback failed: {e}")
        buffer = io.BytesIO()
        buffer.write(f"%PDF-1.4\n1 0 obj << /Title ({title}) >>\n{markdown_content}".encode('utf-8'))
        return buffer.getvalue()

